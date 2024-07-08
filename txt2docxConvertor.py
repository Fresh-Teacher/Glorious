import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage
import re

def add_underline(run):
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    run._element.get_or_add_rPr().append(u)

def create_docx(txt_path, docx_path, header_image_path):
    with open(txt_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    content = content.replace("Not answered", "____________________________________")
    paragraphs = content.split('\n\n')
    
    doc = Document()
    
    pil_img = PILImage.open(header_image_path)
    width, height = pil_img.size
    new_width = Inches(1.5)
    aspect_ratio = height / width
    new_height = new_width * aspect_ratio
    doc.add_picture(header_image_path, width=new_width, height=new_height)
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for para in paragraphs:
        if para.strip():
            if any(header in para for header in ["GLORIOUS KINDERGARTEN & PRIMARY SCHOOL", "MID-TERM II EXAMINATION 2024", "INFORMATION AND COMMUNICATION TECHNOLOGY (ICT)"]):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(para.strip())
                run.bold = True
                run.font.size = Pt(14)
            elif any(field in para for field in ["NAME:", "CLASS:", "STREAM:"]):
                p = doc.add_paragraph()
                run = p.add_run(para.strip())
                run.bold = True
            else:
                lines = para.split('\n')
                current_question = None
                for line in lines:
                    line = line.strip()
                    if re.match(r'^\d+\.\s', line):
                        if current_question:
                            doc.add_paragraph()  # Add space before new question
                        current_question = doc.add_paragraph()
                        current_question.add_run(line)
                    elif current_question:
                        run = current_question.add_run('\n' + line)
                        add_underline(run)
                    else:
                        p = doc.add_paragraph()
                        run = p.add_run(line)
                        add_underline(run)
    
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("© Glorious Schools ICT Department")
    footer_run.font.name = 'Calibri'
    footer_run.font.size = Pt(11)
    footer_run.italic = True
    footer_run.bold = True
    
    doc.save(docx_path)

def process_folder(folder_path, header_image_path):
    for filename in os.listdir(folder_path):
        if filename.endswith('.txt'):
            txt_path = os.path.join(folder_path, filename)
            docx_path = os.path.join(folder_path, filename[:-4] + '.docx')
            create_docx(txt_path, docx_path, header_image_path)
            print(f"Processed: {filename}")

folder_path = 'texts'
header_image_path = 'schoologo-2-150x150.png'

process_folder(folder_path, header_image_path)
print("All files have been processed and converted to DOCX.")